from pathlib import Path

from app.core.config import get_settings
from app.services.cv.template import ensure_template_shape


class CVExportService:
    def __init__(self) -> None:
        self.settings = get_settings()

    def export_pdf(self, user_id: int, cv_id: int, title: str, cv_json: dict) -> str:
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except Exception as exc:
            raise RuntimeError("reportlab is not installed. Run `pip install reportlab`.") from exc

        output_dir = Path(self.settings.generated_assets_dir) / "pdf" / str(user_id)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"cv_{cv_id}.pdf"

        structured = ensure_template_shape(cv_json)
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=14 * mm,
            rightMargin=14 * mm,
            topMargin=12 * mm,
            bottomMargin=12 * mm,
        )
        styles = getSampleStyleSheet()
        name_style = ParagraphStyle("Name", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=18, spaceAfter=4)
        contact_style = ParagraphStyle("Contact", parent=styles["Normal"], fontName="Helvetica", fontSize=9, textColor=colors.HexColor("#333333"))
        section_style = ParagraphStyle(
            "SectionHeader",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11,
            textTransform="uppercase",
            spaceBefore=10,
            spaceAfter=4,
            textColor=colors.HexColor("#0f172a"),
        )
        section_note_style = ParagraphStyle("SectionNote", parent=styles["Normal"], fontName="Helvetica-Oblique", fontSize=8, textColor=colors.HexColor("#475569"), spaceAfter=3)
        body_style = ParagraphStyle("Body", parent=styles["Normal"], fontName="Helvetica", fontSize=9.5, leading=12)
        bullet_style = ParagraphStyle("Bullet", parent=body_style, leftIndent=10, bulletIndent=2, spaceAfter=1)
        item_heading_style = ParagraphStyle("ItemHeading", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10.3, leading=12)
        item_meta_style = ParagraphStyle("ItemMeta", parent=styles["Normal"], fontName="Helvetica", fontSize=9, textColor=colors.HexColor("#334155"), leading=11)

        story = []

        display_name = str(structured["header"].get("name", title) or title)
        story.append(Paragraph(self._escape(display_name), name_style))
        contact_parts = [
            str(structured["header"].get("email", "")).strip(),
            str(structured["header"].get("phone", "")).strip(),
            str(structured["header"].get("location", "")).strip(),
            str(structured["header"].get("linkedin", "")).strip(),
            str(structured["header"].get("github", "")).strip(),
        ]
        contact_line = " | ".join([part for part in contact_parts if part])
        if contact_line:
            story.append(Paragraph(self._escape(contact_line), contact_style))
        story.append(Spacer(1, 4))
        story.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#1f2937")))
        story.append(Spacer(1, 6))

        section_labels = structured.get("section_labels", {})
        section_descriptions = structured.get("section_descriptions", {})
        for section in structured["section_order"]:
            content = structured["sections"].get(section)
            if content in (None, "", []):
                continue

            label = section_labels.get(section, section.replace("_", " ").title())
            story.append(Paragraph(self._escape(label), section_style))
            if section in section_descriptions and section_descriptions[section]:
                story.append(Paragraph(self._escape(str(section_descriptions[section])), section_note_style))
            story.append(HRFlowable(width="100%", thickness=0.4, color=colors.HexColor("#CBD5E1")))
            story.append(Spacer(1, 4))

            if isinstance(content, str):
                story.append(Paragraph(self._escape(content), body_style))
                story.append(Spacer(1, 6))
                continue

            if isinstance(content, dict):
                self._append_structured_item(content, story, item_heading_style, item_meta_style, body_style, bullet_style, Table, TableStyle, colors)
                story.append(Spacer(1, 6))
                continue

            if isinstance(content, list):
                for item in content:
                    if isinstance(item, dict):
                        self._append_structured_item(item, story, item_heading_style, item_meta_style, body_style, bullet_style, Table, TableStyle, colors)
                    else:
                        story.append(Paragraph(self._escape(str(item)), bullet_style, bulletText="\u2022"))
                story.append(Spacer(1, 6))

        doc.build(story)
        return str(output_path)

    def export_docx(self, user_id: int, cv_id: int, title: str, cv_json: dict) -> str:
        try:
            from docx import Document
        except Exception as exc:
            raise RuntimeError("python-docx is not installed. Run `pip install python-docx`.") from exc

        output_dir = Path(self.settings.generated_assets_dir) / "docx" / str(user_id)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"cv_{cv_id}.docx"

        structured = ensure_template_shape(cv_json)
        doc = Document()
        doc.add_heading(structured["header"].get("name") or title, level=1)
        contact_line = " | ".join(
            [
                str(structured["header"].get("email", "")).strip(),
                str(structured["header"].get("phone", "")).strip(),
                str(structured["header"].get("location", "")).strip(),
                str(structured["header"].get("linkedin", "")).strip(),
                str(structured["header"].get("github", "")).strip(),
            ]
        ).strip(" |")
        if contact_line:
            doc.add_paragraph(contact_line)

        for section in structured["section_order"]:
            content = structured["sections"].get(section)
            if content in (None, "", []):
                continue
            label = structured.get("section_labels", {}).get(section, section.replace("_", " ").title())
            doc.add_heading(label, level=2)
            section_note = structured.get("section_descriptions", {}).get(section)
            if section_note:
                doc.add_paragraph(str(section_note))
            if isinstance(content, list):
                for item in content:
                    if isinstance(item, dict):
                        heading = item.get("heading") or item.get("title") or item.get("organization") or ""
                        subheading = item.get("subheading") or item.get("role") or item.get("location") or ""
                        date = item.get("date") or item.get("duration") or ""
                        line = " | ".join([part for part in [heading, subheading, date] if part])
                        if line:
                            doc.add_paragraph(line)
                        description = item.get("description")
                        if description:
                            doc.add_paragraph(str(description))
                        highlights = item.get("highlights", [])
                        for hl in highlights:
                            doc.add_paragraph(str(hl), style="List Bullet")
                    else:
                        doc.add_paragraph(str(item), style="List Bullet")
            elif isinstance(content, str):
                doc.add_paragraph(content)
            else:
                doc.add_paragraph(str(content))
        doc.save(str(output_path))
        return str(output_path)

    @staticmethod
    def _escape(value: str) -> str:
        return (
            value.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )

    def _append_structured_item(
        self,
        item: dict,
        story: list,
        item_heading_style,
        item_meta_style,
        body_style,
        bullet_style,
        table_cls,
        table_style_cls,
        colors,
    ) -> None:
        heading = str(item.get("heading") or item.get("title") or item.get("organization") or "").strip()
        right_meta = str(item.get("date") or item.get("duration") or "").strip()
        if heading or right_meta:
            row = [
                [Paragraph(self._escape(heading), item_heading_style), Paragraph(self._escape(right_meta), item_meta_style)]
            ]
            table = table_cls(row, colWidths=["73%", "27%"])
            table.setStyle(
                table_style_cls(
                    [
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ("TOPPADDING", (0, 0), (-1, -1), 0),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                    ]
                )
            )
            story.append(table)

        secondary_parts = [
            str(item.get("subheading") or item.get("role") or "").strip(),
            str(item.get("location") or "").strip(),
        ]
        secondary_line = " | ".join([part for part in secondary_parts if part])
        if secondary_line:
            story.append(Paragraph(self._escape(secondary_line), item_meta_style))

        description = item.get("description")
        if description:
            story.append(Paragraph(self._escape(str(description)), body_style))

        highlights = item.get("highlights", [])
        if isinstance(highlights, list):
            for hl in highlights:
                story.append(Paragraph(self._escape(str(hl)), bullet_style, bulletText="\u2022"))

        story.append(Spacer(1, 3))
